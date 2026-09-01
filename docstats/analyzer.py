"""Word document analysis: tables, figures, cells, merged cells, nested
tables and tables that break across pages.

Two engines complement each other:

* XML (python-docx) - reads the OOXML directly: table count, rows/columns/
  cells, gridSpan (horizontal merge), vMerge (vertical merge), nested tables,
  figures. Fast, and does not need Word installed.
* Word COM (pywin32) - opens the file with MS Word to get REAL PAGE NUMBERS
  and detect tables that break across pages. A .docx file does not store
  pagination, so this is the only way to get that right.
"""

from __future__ import annotations

import os
import re
import shutil
import tempfile
import traceback
from dataclasses import dataclass, field
from typing import Callable, Iterable, Iterator

from docx import Document
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A = "http://schemas.openxmlformats.org/drawingml/2006/main"
PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
CHART = "http://schemas.openxmlformats.org/drawingml/2006/chart"
DIAGRAM = "http://schemas.openxmlformats.org/drawingml/2006/diagram"
WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
WPG = "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
WPC = "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
VML = "urn:schemas-microsoft-com:vml"

EMU_PER_CM = 360000
TWIP_PER_EMU = 635  # 1 twip = 1/1440 inch, 1 inch = 914400 EMU
A4_TWIPS = (11906, 16838)  # fallback page size when the file declares none

WORD_EXTENSIONS = (".docx", ".docm", ".doc")
PDF_EXTENSIONS = (".pdf",)
DOC_EXTENSIONS = WORD_EXTENSIONS + PDF_EXTENSIONS

#: default threshold: a table with >= this many cells counts as "complex"
DEFAULT_COMPLEX_THRESHOLD = 50

#: default threshold: a figure must cover at least this share of the page to
#: be counted. 0.25 = a quarter of the page, enough to drop logos / stamps /
#: icons. Set to 0 to count every figure.
DEFAULT_FIGURE_MIN_AREA = 0.25

# Caption prefixes, English and Vietnamese - documents in the wild are often
# bilingual. Used to count captions and to attach one to each table.
RE_FIG_CAPTION = re.compile(
    r"^\s*(hình|figure|fig\.?|ảnh|biểu\s*đồ|sơ\s*đồ|chart|image)\s*[:.\-]?\s*\d",
    re.IGNORECASE,
)
RE_TBL_CAPTION = re.compile(
    r"^\s*(bảng|table|biểu)\s*[:.\-]?\s*\d",
    re.IGNORECASE,
)

# Word constants
WD_ACTIVE_END_PAGE = 3  # wdActiveEndPageNumber (absolute, from the file start)
WD_ACTIVE_END_ADJ_PAGE = 1  # wdActiveEndAdjustedPageNumber (printed page number)
WD_STAT_PAGES = 2
WD_FORMAT_DOCX = 16

#: above this many cells, skip the per-row scan for page breaks (too slow)
ROW_SCAN_CELL_LIMIT = 2500


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------


@dataclass
class TableInfo:
    """Everything measured about a single table."""

    file: str
    file_name: str
    index: int  # 1-based top-level table number; nested tables reuse it
    label: str  # "Table 3", or "Table 3.1" for a nested table
    level: int  # 1 = top-level table, >= 2 = nested inside another table
    parent_label: str = ""

    rows: int = 0
    grid_cols: int = 0  # number of grid columns ("width")
    grid_cells: int = 0  # rows x grid_cols = cell count if NOTHING were merged
    cells: int = 0  # VISIBLE cells (a merged block counts as 1) -> COMPLEXITY SCORE
    merged_cells: int = 0  # = grid_cells - cells: CELLS MERGED AWAY
    physical_cells: int = 0  # number of cell tags in the XML (Word files only)
    merged_h: int = 0  # horizontal merge regions
    merged_v: int = 0  # vertical merge regions
    nested_direct: int = 0  # direct child tables
    nested_total: int = 0  # child tables at any depth

    score: int = 0
    is_complex: bool = False
    has_merge: bool = False
    is_nested: bool = False  # this table sits inside another one

    page_start: int | None = None
    page_end: int | None = None
    page_label: str = ""  # printed page number when it differs from the absolute one
    pages_spanned: int = 0
    is_split: bool = False
    split_rows: list[int] = field(default_factory=list)  # page break after row N
    page_is_approx: bool = False

    repeat_header: bool = False  # "repeat header row" is enabled
    rows_cant_split: int = 0  # rows marked "cannot split across pages"
    caption: str = ""
    preview: str = ""

    word_index: int = 0  # index in Word's doc.Tables (0 = unknown)

    # ---- flags that make reviewing easier ----
    @property
    def split_without_header(self) -> bool:
        return self.is_split and not self.repeat_header

    def issue_tags(self) -> list[str]:
        tags = []
        if self.is_complex:
            tags.append("Complex")
        if self.has_merge:
            tags.append("Merged cells")
        if self.nested_direct:
            tags.append("Has nested table")
        if self.is_nested:
            tags.append("Nested table")
        if self.is_split:
            tags.append("Split across pages")
        if self.split_without_header:
            tags.append("Split without header")
        return tags


@dataclass
class FigureInfo:
    """Everything measured about a single figure / image / chart."""

    file: str
    file_name: str
    index: int
    kind: str  # Image / Chart / SmartArt / Shape group / Image (VML) ...
    in_table: bool = False
    page: int | None = None
    caption: str = ""
    width_cm: float = 0.0
    height_cm: float = 0.0
    #: figure area / page area (0 = size unknown)
    area_ratio: float = 0.0
    #: whether it counts towards the figure total (see recompute_figures)
    counted: bool = True

    @property
    def size_known(self) -> bool:
        return self.area_ratio > 0

    @property
    def size_text(self) -> str:
        if not self.size_known:
            return "?"
        return f"{self.width_cm:.1f}×{self.height_cm:.1f}"

    @property
    def area_percent(self) -> float:
        return self.area_ratio * 100


@dataclass
class DocResult:
    """Result for a single document."""

    path: str
    name: str
    file_type: str = "Word"  # "Word" or "PDF"
    size: int = 0
    pages: int | None = None
    scanned_pages: int = 0
    tables: list[TableInfo] = field(default_factory=list)
    figures: list[FigureInfo] = field(default_factory=list)
    fig_captions: int = 0
    tbl_captions: int = 0
    error: str = ""
    word_used: bool = False
    warnings: list[str] = field(default_factory=list)

    # --- aggregates ---
    @property
    def n_tables_top(self) -> int:
        return sum(1 for t in self.tables if t.level == 1)

    @property
    def n_tables_all(self) -> int:
        return len(self.tables)

    @property
    def n_figures(self) -> int:
        """Figures that count (small ones below the area threshold are dropped)."""
        return sum(1 for f in self.figures if f.counted)

    @property
    def n_figures_all(self) -> int:
        """All figures found, including the small ones that were dropped."""
        return len(self.figures)

    @property
    def n_figures_small(self) -> int:
        return sum(1 for f in self.figures if not f.counted)

    @property
    def n_complex(self) -> int:
        return sum(1 for t in self.tables if t.is_complex)

    @property
    def n_split(self) -> int:
        return sum(1 for t in self.tables if t.is_split)

    @property
    def n_nested_parent(self) -> int:
        return sum(1 for t in self.tables if t.nested_direct > 0)

    @property
    def n_nested_child(self) -> int:
        return sum(1 for t in self.tables if t.is_nested)

    @property
    def n_merged(self) -> int:
        return sum(1 for t in self.tables if t.has_merge)

    @property
    def n_split_no_header(self) -> int:
        return sum(1 for t in self.tables if t.split_without_header)

    @property
    def total_cells(self) -> int:
        return sum(t.cells for t in self.tables)

    @property
    def total_merged_cells(self) -> int:
        return sum(t.merged_cells for t in self.tables)

    @property
    def max_score(self) -> int:
        return max((t.score for t in self.tables), default=0)


@dataclass
class ScanResult:
    """Result of a whole scan."""

    root: str = ""
    docs: list[DocResult] = field(default_factory=list)
    complex_threshold: int = DEFAULT_COMPLEX_THRESHOLD
    figure_min_area: float = DEFAULT_FIGURE_MIN_AREA
    word_engine: bool = False
    elapsed: float = 0.0

    @property
    def ok_docs(self) -> list[DocResult]:
        return [d for d in self.docs if not d.error]

    @property
    def failed_docs(self) -> list[DocResult]:
        return [d for d in self.docs if d.error]

    def all_tables(self) -> list[TableInfo]:
        return [t for d in self.docs for t in d.tables]

    def all_figures(self, counted_only: bool = False) -> list[FigureInfo]:
        return [
            f for d in self.docs for f in d.figures if f.counted or not counted_only
        ]

    def totals(self) -> dict[str, int]:
        tables = self.all_tables()
        figures = self.all_figures()
        return {
            "files": len(self.docs),
            "files_word": sum(1 for d in self.docs if d.file_type == "Word"),
            "files_pdf": sum(1 for d in self.docs if d.file_type == "PDF"),
            "files_error": len(self.failed_docs),
            "pages": sum(d.pages or 0 for d in self.docs),
            "tables": len(tables),
            "tables_top": sum(1 for t in tables if t.level == 1),
            "figures": sum(1 for f in figures if f.counted),
            "figures_all": len(figures),
            "figures_small": sum(1 for f in figures if not f.counted),
            "complex": sum(1 for t in tables if t.is_complex),
            "split": sum(1 for t in tables if t.is_split),
            "nested_parent": sum(1 for t in tables if t.nested_direct > 0),
            "nested_child": sum(1 for t in tables if t.is_nested),
            "merged": sum(1 for t in tables if t.has_merge),
            "split_no_header": sum(1 for t in tables if t.split_without_header),
            "cells": sum(t.cells for t in tables),
            "merged_cells": sum(t.merged_cells for t in tables),
        }


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------


def _tag(ns: str, name: str) -> str:
    return f"{{{ns}}}{name}"


def _text_of(el) -> str:
    return "".join(t.text or "" for t in el.iter(qn("w:t")))


def _int_attr(el, name: str, default: int = 0) -> int:
    if el is None:
        return default
    val = el.get(qn(name))
    try:
        return int(val)
    except (TypeError, ValueError):
        return default


def _direct(el, ns: str, name: str) -> list:
    return el.findall(_tag(ns, name))


def _rows_of(tbl) -> list:
    """Direct w:tr children, looking through a wrapping w:sdt if present."""
    rows = _direct(tbl, W, "tr")
    for sdt in _direct(tbl, W, "sdt"):
        for content in sdt.findall(_tag(W, "sdtContent")):
            rows.extend(_direct(content, W, "tr"))
    return rows


def _cells_of(tr) -> list:
    cells = _direct(tr, W, "tc")
    for sdt in _direct(tr, W, "sdt"):
        for content in sdt.findall(_tag(W, "sdtContent")):
            cells.extend(_direct(content, W, "tc"))
    return cells


def _nested_tables_of_cell(tc) -> list:
    """Direct child tables of one cell."""
    tables = _direct(tc, W, "tbl")
    for sdt in _direct(tc, W, "sdt"):
        for content in sdt.findall(_tag(W, "sdtContent")):
            tables.extend(_direct(content, W, "tbl"))
    return tables


# ---------------------------------------------------------------------------
# Single table analysis (XML)
# ---------------------------------------------------------------------------


def _finalize(info: TableInfo, threshold: int) -> None:
    """Derive the computed fields; shared by the Word and the PDF engine.

    Ground rule: a table with nothing merged has exactly ``rows x columns``
    cells. However many cells are missing from that number is exactly how
    many cells were **merged away**.
    """
    info.grid_cells = info.rows * info.grid_cols
    info.merged_cells = max(0, info.grid_cells - info.cells)
    info.has_merge = info.merged_cells > 0
    info.score = info.cells
    info.is_complex = info.cells >= threshold


def _analyze_table(
    tbl,
    *,
    doc_path: str,
    doc_name: str,
    index: int,
    label: str,
    level: int,
    parent_label: str,
    threshold: int,
    out: list[TableInfo],
) -> TableInfo:
    """Analyze one table recursively, appending nested tables to ``out``."""
    info = TableInfo(
        file=doc_path,
        file_name=doc_name,
        index=index,
        label=label,
        level=level,
        parent_label=parent_label,
        is_nested=level > 1,
    )
    out.append(info)

    grid = tbl.find(_tag(W, "tblGrid"))
    declared_cols = len(_direct(grid, W, "gridCol")) if grid is not None else 0

    rows = _rows_of(tbl)
    info.rows = len(rows)

    preview_bits: list[str] = []
    child_specs: list[tuple] = []
    max_width = 0

    # track every column so vertical merge regions are detected even when the
    # file omits vMerge="restart" (many docx generators do)
    active_v: set[int] = set()  # columns with an open vertical merge
    prev_cols: set[int] = set()  # columns that had a cell on the row above

    for r_i, tr in enumerate(rows):
        tr_pr = tr.find(_tag(W, "trPr"))
        if tr_pr is not None:
            if tr_pr.find(_tag(W, "tblHeader")) is not None:
                info.repeat_header = True
            if tr_pr.find(_tag(W, "cantSplit")) is not None:
                info.rows_cant_split += 1

        cells = _cells_of(tr)
        col = 0
        cur_cols: set[int] = set()
        cur_active: set[int] = set()

        for c_i, tc in enumerate(cells):
            info.physical_cells += 1
            tc_pr = tc.find(_tag(W, "tcPr"))
            span = 1
            v_state = None
            if tc_pr is not None:
                span = max(1, _int_attr(tc_pr.find(_tag(W, "gridSpan")), "w:val", 1))
                v_merge = tc_pr.find(_tag(W, "vMerge"))
                if v_merge is not None:
                    v_state = (v_merge.get(qn("w:val")) or "continue").lower()

            if span > 1:
                info.merged_h += 1

            if v_state == "continue" and (col in active_v or col in prev_cols):
                # swallowed by a vertical merge above -> not a visible cell
                if col not in active_v:
                    info.merged_v += 1  # the merge region starts one row up
                cur_active.add(col)
            else:
                info.cells += 1
                if v_state == "restart":
                    info.merged_v += 1
                    cur_active.add(col)

            if r_i == 0 and len(preview_bits) < 6:
                txt = " ".join(_text_of(tc).split())
                if txt:
                    preview_bits.append(txt[:24])

            for sub in _nested_tables_of_cell(tc):
                info.nested_direct += 1
                child_specs.append((sub, r_i + 1, c_i + 1))

            cur_cols.add(col)
            col += span

        max_width = max(max_width, col)
        active_v, prev_cols = cur_active, cur_cols

    info.grid_cols = max_width or declared_cols
    _finalize(info, threshold)
    info.preview = " | ".join(preview_bits)

    # nested tables
    for n_i, (sub, r_i, c_i) in enumerate(child_specs, start=1):
        child = _analyze_table(
            sub,
            doc_path=doc_path,
            doc_name=doc_name,
            index=index,
            label=f"{label}.{n_i}",
            level=level + 1,
            parent_label=label,
            threshold=threshold,
            out=out,
        )
        child.preview = f"[inside {label} cell R{r_i}C{c_i}] " + child.preview
        info.nested_total += 1 + child.nested_total

    return info


# ---------------------------------------------------------------------------
# Figures
# ---------------------------------------------------------------------------


def _classify_drawing(el) -> str | None:
    """Classify a w:drawing element. None = skip it (plain text box)."""
    if el.find(f".//{_tag(PIC, 'pic')}") is not None or el.find(f".//{_tag(A, 'blip')}") is not None:
        return "Image"
    if el.find(f".//{_tag(CHART, 'chart')}") is not None:
        return "Chart"
    if el.find(f".//{_tag(DIAGRAM, 'relIds')}") is not None:
        return "SmartArt"
    if el.find(f".//{_tag(WPC, 'wpc')}") is not None:
        return "Drawing canvas"
    if el.find(f".//{_tag(WPG, 'wgp')}") is not None or el.find(f".//{_tag(WPG, 'grpSp')}") is not None:
        return "Shape group"
    if el.find(f".//{_tag(WPS, 'wsp')}") is not None:
        return None  # standalone text box / shape -> not a figure
    return "Graphic object"


def _classify_pict(el) -> str | None:
    if el.find(f".//{_tag(VML, 'imagedata')}") is not None:
        return "Image (VML)"
    if el.find(f".//{_tag(VML, 'textbox')}") is not None:
        return None
    return "Shape (VML)"


RE_VML_LEN = re.compile(r"(-?[\d.]+)\s*(pt|in|cm|mm|px|pc)?", re.IGNORECASE)
VML_UNIT_EMU = {  # one unit expressed in EMU
    "pt": 12700,
    "in": 914400,
    "cm": 360000,
    "mm": 36000,
    "px": 9525,  # 96 dpi
    "pc": 152400,
    "": 12700,  # VML defaults to points
}


def _page_area_emu(body) -> float:
    """Area of one page (EMU2), taken from the first section's page size."""
    w_tw, h_tw = A4_TWIPS
    pg = body.find(f".//{_tag(W, 'sectPr')}/{_tag(W, 'pgSz')}")
    if pg is not None:
        w_tw = _int_attr(pg, "w:w", w_tw) or w_tw
        h_tw = _int_attr(pg, "w:h", h_tw) or h_tw
    return (w_tw * TWIP_PER_EMU) * (h_tw * TWIP_PER_EMU)


def _plain_int(el, name: str, default: int = 0) -> int:
    """Read an attribute with NO namespace (e.g. cx/cy of wp:extent)."""
    try:
        return int(float(el.get(name)))
    except (TypeError, ValueError):
        return default


def _vml_length_emu(text: str) -> int:
    m = RE_VML_LEN.match(text.strip())
    if not m:
        return 0
    try:
        value = float(m.group(1))
    except ValueError:
        return 0
    return int(abs(value) * VML_UNIT_EMU.get((m.group(2) or "").lower(), 12700))


def _figure_size_emu(el) -> tuple[int, int]:
    """Rendered size of one figure, in EMU. (0, 0) = could not measure."""
    # DrawingML: wp:extent is the size actually rendered on the page
    ext = el.find(f".//{_tag(WP, 'extent')}")
    if ext is None:
        ext = el.find(f".//{_tag(A, 'ext')}")  # fallback: a:ext inside xfrm
    if ext is not None:
        # cx/cy are bare attributes, no namespace -> _int_attr/qn do not apply
        cx, cy = _plain_int(ext, "cx"), _plain_int(ext, "cy")
        if cx > 0 and cy > 0:
            return cx, cy

    # VML (w:pict): the size lives in style="width:36pt;height:36pt"
    for shape in el.iter():
        style = shape.get("style")
        if not style:
            continue
        dims = {}
        for part in style.split(";"):
            key, _, value = part.partition(":")
            key = key.strip().lower()
            if key in ("width", "height") and value:
                dims[key] = _vml_length_emu(value)
        if dims.get("width") and dims.get("height"):
            return dims["width"], dims["height"]
    return 0, 0


def _collect_figures(body, doc_path: str, doc_name: str) -> list[FigureInfo]:
    figures: list[FigureInfo] = []
    page_area = _page_area_emu(body)
    for el in body.iter():
        kind = None
        if el.tag == _tag(W, "drawing"):
            kind = _classify_drawing(el)
        elif el.tag == _tag(W, "pict"):
            # a w:pict inside a w:object is already wrapped -> still counted once
            kind = _classify_pict(el)
        if not kind:
            continue
        in_table = False
        parent = el.getparent()
        while parent is not None:
            if parent.tag == _tag(W, "tbl"):
                in_table = True
                break
            parent = parent.getparent()
        cx, cy = _figure_size_emu(el)
        figures.append(
            FigureInfo(
                file=doc_path,
                file_name=doc_name,
                index=len(figures) + 1,
                kind=kind,
                in_table=in_table,
                width_cm=cx / EMU_PER_CM,
                height_cm=cy / EMU_PER_CM,
                area_ratio=(cx * cy) / page_area if cx and cy and page_area else 0.0,
            )
        )
    return figures


# ---------------------------------------------------------------------------
# Single file analysis with python-docx
# ---------------------------------------------------------------------------


def _analyze_xml(path: str, display_path: str, threshold: int) -> DocResult:
    res = DocResult(path=display_path, name=os.path.basename(display_path))
    try:
        res.size = os.path.getsize(display_path)
    except OSError:
        pass

    document = Document(path)
    body = document.element.body

    tables: list[TableInfo] = []
    top_index = 0
    for tbl in _direct(body, W, "tbl"):
        top_index += 1
        _analyze_table(
            tbl,
            doc_path=display_path,
            doc_name=res.name,
            index=top_index,
            label=f"Table {top_index}",
            level=1,
            parent_label="",
            threshold=threshold,
            out=tables,
        )
    # tables wrapped in a content control at body level
    for sdt in _direct(body, W, "sdt"):
        for content in sdt.findall(_tag(W, "sdtContent")):
            for tbl in _direct(content, W, "tbl"):
                top_index += 1
                _analyze_table(
                    tbl,
                    doc_path=display_path,
                    doc_name=res.name,
                    index=top_index,
                    label=f"Table {top_index}",
                    level=1,
                    parent_label="",
                    threshold=threshold,
                    out=tables,
                )

    # number top-level tables the way Word's doc.Tables collection does
    w_idx = 0
    for t in tables:
        if t.level == 1:
            w_idx += 1
            t.word_index = w_idx

    res.tables = tables
    res.figures = _collect_figures(body, display_path, res.name)

    # count captions and attach one per table (paragraph just above / below)
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if RE_FIG_CAPTION.match(text):
            res.fig_captions += 1
        elif RE_TBL_CAPTION.match(text):
            res.tbl_captions += 1

    _attach_captions(body, res)
    return res


def _attach_captions(body, res: DocResult) -> None:
    """Find the caption ('Table 3: ...') right before/after each top-level table."""
    children = list(body)
    top_tables = [t for t in res.tables if t.level == 1]
    tbl_positions = [i for i, el in enumerate(children) if el.tag == _tag(W, "tbl")]
    for t, pos in zip(top_tables, tbl_positions):
        for probe in (pos - 1, pos + 1, pos - 2):
            if 0 <= probe < len(children) and children[probe].tag == _tag(W, "p"):
                text = " ".join(_text_of(children[probe]).split())
                if text and RE_TBL_CAPTION.match(text):
                    t.caption = text[:120]
                    break


# ---------------------------------------------------------------------------
# Word COM engine
# ---------------------------------------------------------------------------


class WordEngine:
    """Drives MS Word over COM to obtain real page numbers."""

    def __init__(self) -> None:
        import pythoncom  # noqa: F401  (make sure pywin32 is available)
        import win32com.client

        pythoncom.CoInitialize()
        self._pythoncom = pythoncom
        self.app = win32com.client.DispatchEx("Word.Application")
        self.app.Visible = False
        self.app.DisplayAlerts = 0
        try:
            self.app.ScreenUpdating = False
            self.app.Options.Pagination = True
        except Exception:
            pass
        self._tmpdir = tempfile.mkdtemp(prefix="docstats_")

    # -- lifecycle -------------------------------------------------------
    def close(self) -> None:
        try:
            self.app.Quit(0)
        except Exception:
            pass
        try:
            self._pythoncom.CoUninitialize()
        except Exception:
            pass
        shutil.rmtree(self._tmpdir, ignore_errors=True)

    def __enter__(self) -> "WordEngine":
        return self

    def __exit__(self, *exc) -> None:
        self.close()

    # -- helpers ---------------------------------------------------------
    def _open(self, path: str):
        return self.app.Documents.Open(
            os.path.abspath(path),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
            PasswordDocument="#docstats#",  # keep Word from hanging on a password prompt
            Visible=False,
        )

    @staticmethod
    def _page_at(doc, pos: int) -> int:
        rng = doc.Range(pos, pos)
        return int(rng.Information(WD_ACTIVE_END_PAGE))

    @staticmethod
    def _page_label_at(doc, pos: int) -> int:
        rng = doc.Range(pos, pos)
        return int(rng.Information(WD_ACTIVE_END_ADJ_PAGE))

    def convert_to_docx(self, path: str) -> str:
        """.doc -> a temporary .docx so python-docx can read it."""
        doc = self._open(path)
        try:
            target = os.path.join(
                self._tmpdir, f"{abs(hash(path))}_{os.path.basename(path)}x"
            )
            doc.SaveAs2(target, FileFormat=WD_FORMAT_DOCX)
            return target
        finally:
            doc.Close(0)

    # -- page numbers ----------------------------------------------------
    def page_info(self, path: str, res: DocResult) -> None:
        doc = self._open(path)
        try:
            try:
                doc.Repaginate()
            except Exception:
                pass
            try:
                res.pages = int(doc.ComputeStatistics(WD_STAT_PAGES))
            except Exception:
                res.pages = None

            top_tables = [t for t in res.tables if t.level == 1]
            word_tables = doc.Tables
            if word_tables.Count != len(top_tables):
                res.warnings.append(
                    f"Word counted {word_tables.Count} top-level tables but the "
                    f"XML counted {len(top_tables)} - page numbers may be off."
                )
            n = min(word_tables.Count, len(top_tables))
            for i in range(1, n + 1):
                self._fill_table_pages(doc, word_tables.Item(i), top_tables[i - 1], res)

            self._fill_figure_pages(doc, res)
            res.word_used = True
        finally:
            doc.Close(0)

    def _fill_table_pages(self, doc, w_table, info: TableInfo, res: DocResult) -> None:
        try:
            rng = w_table.Range
            start, end = int(rng.Start), int(rng.End)
            info.page_start = self._page_at(doc, start)
            info.page_end = self._page_at(doc, max(start, end - 1))
            label_start = self._page_label_at(doc, start)
            if label_start != info.page_start:
                info.page_label = str(label_start)
            info.pages_spanned = (info.page_end or 0) - (info.page_start or 0) + 1
            info.is_split = info.pages_spanned > 1
            if info.is_split:
                info.split_rows = self._split_rows(doc, w_table, info)
        except Exception as exc:  # pragma: no cover - environment dependent
            res.warnings.append(f"Could not read the page of {info.label}: {exc}")
            return

        # nested tables: take the page from Word's own nested collection when
        # the counts line up, otherwise inherit the parent and flag it approximate
        children = [t for t in res.tables if t.parent_label == info.label]
        if not children:
            return
        try:
            nested = w_table.Tables
            if nested.Count == len(children):
                for w_child, child in zip(
                    (nested.Item(i) for i in range(1, nested.Count + 1)), children
                ):
                    self._fill_table_pages(doc, w_child, child, res)
                return
        except Exception:
            pass
        for child in children:
            child.page_start = info.page_start
            child.page_end = info.page_end
            child.page_label = info.page_label
            child.pages_spanned = info.pages_spanned
            child.page_is_approx = True

    def _split_rows(self, doc, w_table, info: TableInfo) -> list[int]:
        """Return the list of rows a page break falls after."""
        if info.cells > ROW_SCAN_CELL_LIMIT:
            return []
        first_start: dict[int, int] = {}
        try:
            for cell in w_table.Range.Cells:
                r_i = int(cell.RowIndex)
                pos = int(cell.Range.Start)
                if r_i not in first_start or pos < first_start[r_i]:
                    first_start[r_i] = pos
        except Exception:
            return []
        breaks: list[int] = []
        prev_page = None
        for r_i in sorted(first_start):
            page = self._page_at(doc, first_start[r_i])
            if prev_page is not None and page > prev_page:
                breaks.append(r_i - 1)
            prev_page = page
        return breaks

    def _fill_figure_pages(self, doc, res: DocResult) -> None:
        pages: list[int] = []
        try:
            for i in range(1, doc.InlineShapes.Count + 1):
                shape = doc.InlineShapes.Item(i)
                pages.append(int(shape.Range.Information(WD_ACTIVE_END_PAGE)))
        except Exception:
            pass
        try:
            for i in range(1, doc.Shapes.Count + 1):
                shape = doc.Shapes.Item(i)
                try:
                    pages.append(int(shape.Anchor.Information(WD_ACTIVE_END_PAGE)))
                except Exception:
                    continue
        except Exception:
            pass
        pages.sort()
        if len(pages) == len(res.figures):
            for fig, page in zip(res.figures, pages):
                fig.page = page
        elif pages:
            res.warnings.append(
                f"Word sees {len(pages)} graphic objects, the XML sees "
                f"{len(res.figures)} - figure page numbers are indicative only."
            )
            for fig, page in zip(res.figures, pages):
                fig.page = page


def open_in_word(path: str, word_index: int = 0, page: int | None = None) -> bool:
    """Open the file in Word (a real window) and jump to the table / page.

    Uses ``Dispatch`` so an already-open Word window is reused.
    Returns False when COM is unavailable, so the caller can fall back.
    """
    try:
        import pythoncom
        import win32com.client
    except ImportError:
        return False

    pythoncom.CoInitialize()
    try:
        app = win32com.client.Dispatch("Word.Application")
        app.Visible = True
        doc = None
        target = os.path.normcase(os.path.abspath(path))
        for i in range(1, app.Documents.Count + 1):
            opened = app.Documents.Item(i)
            if os.path.normcase(opened.FullName) == target:
                doc = opened
                break
        if doc is None:
            doc = app.Documents.Open(os.path.abspath(path), AddToRecentFiles=False)
        doc.Activate()
        try:
            if word_index and word_index <= doc.Tables.Count:
                doc.Tables.Item(word_index).Range.Select()
            elif page:
                app.Selection.GoTo(What=1, Which=1, Count=page)  # wdGoToPage
        except Exception:
            pass
        try:
            app.ActiveWindow.Activate()
        except Exception:
            pass
        return True
    except Exception:
        return False
    finally:
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def open_pdf_at_page(path: str, page: int | None = None) -> bool:
    """Open a PDF at a page via ``#page=N`` (Edge/Chrome/Acrobat honour it)."""
    import webbrowser
    from urllib.parse import quote

    if page:
        url = "file:///" + quote(os.path.abspath(path).replace("\\", "/"))
        try:
            if webbrowser.open(f"{url}#page={page}"):
                return True
        except Exception:
            pass
    try:
        os.startfile(os.path.abspath(path))  # noqa: S606
        return True
    except Exception:
        return False


def open_document(path: str, word_index: int = 0, page: int | None = None) -> bool:
    """Open a document at the right table (Word) or the right page (PDF)."""
    if path.lower().endswith(PDF_EXTENSIONS):
        return open_pdf_at_page(path, page)
    return open_in_word(path, word_index, page)


# ---------------------------------------------------------------------------
# File / folder level API
# ---------------------------------------------------------------------------


def find_documents(root: str, recursive: bool = True) -> list[str]:
    """List the Word / PDF files in a folder, skipping ~$ lock files."""
    out: list[str] = []
    if os.path.isfile(root):
        return [root] if root.lower().endswith(DOC_EXTENSIONS) else []
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if not d.startswith((".", "~$"))]
        for name in filenames:
            if name.startswith("~$"):
                continue
            if name.lower().endswith(DOC_EXTENSIONS):
                out.append(os.path.join(dirpath, name))
        if not recursive:
            break
    out.sort(key=lambda p: p.lower())
    return out


def analyze_file(
    path: str,
    threshold: int = DEFAULT_COMPLEX_THRESHOLD,
    engine: WordEngine | None = None,
    pdf_borderless: bool = False,
) -> DocResult:
    """Analyze a single Word or PDF file.

    When ``engine`` is not None, Word files also get real page numbers from
    MS Word. PDFs never need Word.
    """
    name = os.path.basename(path)
    try:
        if path.lower().endswith(PDF_EXTENSIONS):
            from .pdf_analyzer import analyze_pdf

            return analyze_pdf(path, threshold, pdf_borderless)

        source = path
        if path.lower().endswith(".doc"):
            if engine is None:
                res = DocResult(path=path, name=name)
                res.error = "Reading a .doc file requires the Word engine."
                return res
            source = engine.convert_to_docx(path)

        res = _analyze_xml(source, path, threshold)

        if engine is not None:
            try:
                engine.page_info(path, res)
            except Exception as exc:
                res.warnings.append(f"Could not get page numbers from Word: {exc}")
        return res
    except Exception as exc:
        res = DocResult(
            path=path,
            name=name,
            file_type="PDF" if path.lower().endswith(PDF_EXTENSIONS) else "Word",
        )
        try:
            res.size = os.path.getsize(path)
        except OSError:
            pass
        res.error = f"{type(exc).__name__}: {exc}"
        res.warnings.append(traceback.format_exc(limit=3))
        return res


def analyze_folder(
    root: str,
    *,
    recursive: bool = True,
    threshold: int = DEFAULT_COMPLEX_THRESHOLD,
    figure_min_area: float = DEFAULT_FIGURE_MIN_AREA,
    use_word: bool = True,
    pdf_borderless: bool = False,
    progress: Callable[[int, int, str], None] | None = None,
    should_stop: Callable[[], bool] | None = None,
) -> ScanResult:
    """Scan a whole folder.

    ``figure_min_area`` is the minimum share of the page a figure must cover
    to be counted (0.25 = a quarter of the page, 0 = count every figure).
    ``progress(done, total, current_file)`` is called after each file.
    Scanning stops early as soon as ``should_stop()`` returns True.
    """
    import time

    started = time.time()
    files = find_documents(root, recursive)
    has_word_file = any(f.lower().endswith(WORD_EXTENSIONS) for f in files)
    result = ScanResult(
        root=root,
        complex_threshold=threshold,
        figure_min_area=figure_min_area,
        word_engine=use_word and has_word_file,
    )

    engine: WordEngine | None = None
    if use_word and has_word_file:
        try:
            engine = WordEngine()
        except Exception as exc:
            result.word_engine = False
            engine = None
            print(f"[docstats] Could not start Word: {exc}")

    try:
        for i, path in enumerate(files, start=1):
            if should_stop and should_stop():
                break
            if progress:
                progress(i - 1, len(files), path)
            result.docs.append(analyze_file(path, threshold, engine, pdf_borderless))
            if progress:
                progress(i, len(files), path)
    finally:
        if engine is not None:
            engine.close()

    recompute_figures(result, figure_min_area)
    result.elapsed = time.time() - started
    return result


def recompute_complexity(result: ScanResult, threshold: int) -> None:
    """Change the 'complex table' threshold without rescanning."""
    result.complex_threshold = threshold
    for table in result.all_tables():
        table.is_complex = table.cells >= threshold


def recompute_figures(result: ScanResult, min_area: float) -> None:
    """Change the figure area threshold without rescanning.

    ``min_area`` is a share of the page area (0.25 = a quarter, 0 = count
    everything). Figures whose size could not be measured are always counted,
    so nothing is dropped silently.
    """
    result.figure_min_area = min_area
    for fig in result.all_figures():
        fig.counted = min_area <= 0 or not fig.size_known or fig.area_ratio >= min_area
